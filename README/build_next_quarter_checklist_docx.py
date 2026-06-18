from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
OUTPUT_FILE = BASE_DIR / "NEXT_QUARTER_LIFT_REPORT_CHECKLIST.docx"


ACCENT = "1F4E79"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F2F2"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_table(table, header=True):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
            if header and row_index == 0:
                set_cell_shading(cell, ACCENT)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor.from_string(WHITE)
                        run.bold = True


def add_heading(document, text, level=1):
    paragraph = document.add_paragraph()
    paragraph.style = f"Heading {level}"
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor.from_string(ACCENT)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(11)
    return paragraph


def add_note(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(9)
    return paragraph


def add_checklist_table(document, title, items, columns=2):
    add_heading(document, title, 2)
    rows_needed = (len(items) + columns - 1) // columns
    table = document.add_table(rows=rows_needed, cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    index = 0
    for row in table.rows:
        for cell in row.cells:
            if index < len(items):
                set_cell_text(cell, f"[ ] {items[index]}")
            else:
                set_cell_text(cell, "")
                set_cell_shading(cell, LIGHT_GRAY)
            index += 1
    document.add_paragraph()


def add_two_column_table(document, title, headers, rows):
    add_heading(document, title, 2)
    table = document.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=WHITE)
        set_cell_shading(table.rows[0].cells[idx], ACCENT)
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            set_cell_text(row.cells[idx], value)
    style_table(table, header=False)
    document.add_paragraph()


def build_document():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Next Quarter LIFT / eSRE Reporting Checklist")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(ACCENT)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subrun = subtitle.add_run("For LGU Treasurer's Office coordination with Budget and Accounting")
    subrun.font.name = "Arial"
    subrun.font.size = Pt(10)

    meta = document.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    meta_data = [
        ("Year", ""),
        ("Quarter", ""),
        ("Date Prepared", ""),
        ("Prepared By", ""),
    ]
    for row, (label, value) in zip(meta.rows, meta_data):
        set_cell_text(row.cells[0], label, bold=True)
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        set_cell_text(row.cells[1], value)
    document.add_paragraph()

    add_heading(document, "1. Before Encoding", 1)
    add_checklist_table(
        document,
        "Initial Checks",
        [
            "Confirm year and quarter",
            "Confirm report period is open in LIFT",
            "Prepare Budget Office documents",
            "Prepare Accounting Office documents",
            "Prepare Treasurer's Office documents",
            "Back up previous submitted reports",
        ],
    )

    add_heading(document, "2. Source Documents", 1)
    add_checklist_table(
        document,
        "Budget Office",
        [
            "Annual Budget / Appropriation Ordinance",
            "Supplemental Budget Ordinance, if any",
            "Realignment / Augmentation document, if any",
            "Continuing Appropriation list",
            "Budget Appropriation per office",
            "Budget Appropriation per function",
            "Budget Appropriation per PPA",
            "Budget Appropriation per allotment class",
            "Fund classification: GF / SEF / Trust Fund / SHF",
            "PS / MOOE / FE / CO breakdown",
        ],
    )
    add_checklist_table(
        document,
        "Accounting Office",
        [
            "Actual Expenditures per quarter",
            "Disbursement summary",
            "DV summary",
            "Check disbursement summary",
            "Accounts Payable list",
            "Prior-year AP payments",
            "Debt service actual payments",
            "Continuing Appropriation actual expenditures",
            "Office / function / PPA / NGAS code",
            "PS / MOOE / FE / CO classification",
        ],
    )
    add_checklist_table(
        document,
        "Treasurer's Office",
        [
            "Collections / receipts summary",
            "RPT collection",
            "CTC / Cedula collection",
            "Business tax collection",
            "Other fees and charges",
            "Trust fund receipts",
            "Debt service payment confirmation",
            "Beginning cash balance",
            "Ending cash / bank balance",
            "SRE-NGAS reconciliation data",
        ],
    )

    add_heading(document, "3. Firebird Collection Reports To Generate", 1)
    add_two_column_table(
        document,
        "Collection Report Commands",
        ["Report No.", "Purpose / Command"],
        [
            ("List", "python .\\run_collection_query.py --list"),
            ("21", "Summary of Collection: python .\\run_collection_query.py 21 YYYY-MM-DD YYYY-MM-DD"),
            ("22", "Summary of Collection no RPT: python .\\run_collection_query.py 22 YYYY-MM-DD YYYY-MM-DD"),
            ("23", "Summary of Collection RPT: python .\\run_collection_query.py 23 YYYY-MM-DD YYYY-MM-DD"),
            ("25", "Record of Real Property Tax Collection: python .\\run_collection_query.py 25 YYYY-MM-DD YYYY-MM-DD"),
            ("26", "RPT Advance Payment Report: python .\\run_collection_query.py 26 YYYY-MM-DD YYYY-MM-DD"),
            ("27", "Summary Report Sharing: python .\\run_collection_query.py 27 YYYY-MM-DD YYYY-MM-DD"),
        ],
    )

    add_heading(document, "4. LIFT Encoding Areas", 1)
    add_two_column_table(
        document,
        "Where To Encode / Check",
        ["Area", "LIFT Path"],
        [
            ("Budget Appropriation", "Authorized Budget > Budget Appropriation > Expenditures"),
            ("Debt Service Budget", "Authorized Budget > Budget Appropriation > Debt Services"),
            ("Unappropriated Surplus", "Authorized Budget > Budget Appropriation > Unappropriated Surplus"),
            ("RPT Receipts", "Actual Transaction > Receipts > Real Property Tax"),
            ("General Collections", "Actual Transaction > Receipts > General Collections"),
            ("Trust Fund Receipts", "Actual Transaction > Receipts > Trust Fund Receipts"),
            ("Expenditures", "Actual Transaction > Expenditures > Expenditures"),
            ("Accounts Payable", "Actual Transaction > Expenditures > Accounts Payable"),
            ("Debt Services", "Actual Transaction > Expenditures > Debt Services"),
            ("Fund/Cash Balance", "Actual Transaction > Others > Fund/Cash Balance"),
            ("SRE-NGAS Reconciliation", "Actual Transaction > Others > SRE-NGAS Reconciliation"),
        ],
    )

    add_heading(document, "5. Report Checks", 1)
    add_checklist_table(
        document,
        "SRS - Statement of Receipts Sources",
        [
            "RPT Basic agrees with municipal GF share",
            "SEF agrees with municipal SEF share",
            "Business Tax agrees with collection summary",
            "Community Tax agrees with CTC report",
            "Permits and Licenses agree with Other Fees report",
            "Service/User Charges agree with Other Fees report",
            "Economic Enterprises agree with Other Fees report",
            "NTA agrees with official source",
            "Grand Total GF + SEF agrees with SRE receipts",
        ],
    )
    add_checklist_table(
        document,
        "SOE - Statement of Expenditures",
        [
            "Budget Appropriation encoded",
            "Supplemental Budget encoded, if any",
            "Actual Expenditures encoded",
            "Debt Service encoded",
            "Accounts Payable encoded",
            "Continuing Appropriation encoded",
            "PS / MOOE / FE / CO correct",
            "Office / function / PPA correct",
            "GF / SEF fund classification correct",
            "No unexplained negative balance",
        ],
    )
    add_checklist_table(
        document,
        "SRE - Statement of Receipts and Expenditures",
        [
            "SRE receipts agree with SRS grand total",
            "SRE expenditures agree with SOE total expenditures",
            "Beginning cash balance agrees with previous ending balance",
            "Ending fund/cash balance is not negative",
            "Fund/Cash Balance End agrees with LIFT module",
            "NTA agrees with official source",
            "Non-income receipts checked",
            "Non-operating expenditures checked",
            "AP payments checked",
            "Continuing appropriations checked",
        ],
    )

    add_heading(document, "6. If SOE Or Fund/Cash Balance Is Negative", 1)
    add_two_column_table(
        document,
        "Troubleshooting Order",
        ["Step", "Check"],
        [
            ("1", "Identify the exact row with negative value."),
            ("2", "Compare budget appropriation against actual expenditure."),
            ("3", "Check Supplemental Budget, realignment, or augmentation."),
            ("4", "Confirm fund: General Fund, SEF, Trust Fund, or SHF."),
            ("5", "Confirm allotment class: PS, MOOE, FE, or CO."),
            ("6", "Confirm office, function, PPA, and NGAS code."),
            ("7", "Check Accounts Payable and prior-year AP payments."),
            ("8", "Check Continuing Appropriation entries."),
            ("9", "Verify beginning cash balance against prior ending balance."),
            ("10", "Verify ending cash / bank balance and regenerate SRE/SOE."),
        ],
    )
    add_two_column_table(
        document,
        "Who To Ask",
        ["Office", "Ask For"],
        [
            ("Budget Office", "Appropriation, Supplemental Budget, realignment, augmentation, continuing appropriation."),
            ("Accounting Office", "Actual expenditures, disbursements, AP, debt service, and classification."),
            ("Treasurer's Office", "Receipts, cash balance, bank balance, and SRE/SRS/SOE reconciliation."),
        ],
    )

    add_heading(document, "7. Before Submission", 1)
    add_checklist_table(
        document,
        "Final Submission Checks",
        [
            "SRS generated and reviewed",
            "SOE generated and reviewed",
            "SRE generated and reviewed",
            "Alerts checked",
            "Negative values explained or corrected",
            "Budget Office confirmed budget figures",
            "Accounting Office confirmed expenditure figures",
            "Treasurer confirmed receipts and cash balance",
            "Certified correct by authorized officer",
            "Backup/export saved",
        ],
    )
    add_checklist_table(
        document,
        "Files To Save",
        [
            "SRE Excel/PDF",
            "SRS Excel/PDF",
            "SOE Excel/PDF",
            "Supporting collection reports",
            "Budget support files",
            "Accounting support files",
            "Screenshots of resolved alerts, if needed",
        ],
    )

    add_heading(document, "Simple Reminder", 1)
    reminder = document.add_table(rows=3, cols=2)
    reminder.style = "Table Grid"
    reminders = [
        ("Treasurer", "Receipts, cash, fund balance, report checking."),
        ("Budget", "Appropriations, Supplemental Budget, continuing appropriation."),
        ("Accounting", "Expenditures, AP, disbursements, debt service classification."),
    ]
    for row, (office, role) in zip(reminder.rows, reminders):
        set_cell_text(row.cells[0], office, bold=True)
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        set_cell_text(row.cells[1], role)
    document.add_paragraph()
    add_note(document, "If a negative value appears, it is usually a reconciliation issue, not automatically a Treasurer mistake.")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("LGU Treasurer Reporting Checklist - LIFT / eSRE")
    footer_run.font.name = "Arial"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor.from_string("666666")

    document.save(OUTPUT_FILE)
    return OUTPUT_FILE


if __name__ == "__main__":
    print(build_document())
